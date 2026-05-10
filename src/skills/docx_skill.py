"""Generate a clean .docx file from a structured outline."""

from __future__ import annotations

import os
import uuid

from docx import Document
from docx.shared import Pt, RGBColor

OUTPUT_DIR = os.environ.get("FILES_OUTPUT_DIR", "/app/data/files")


def generate_docx(title: str, sections: list[dict]) -> str:
    """Build a .docx and return the absolute file path.

    `sections`: list of dicts. Each dict supports:
        - "heading": str (optional) — section heading
        - "paragraphs": list[str] — paragraphs of plain text
        - "bullets": list[str] (optional) — bullet list
    """
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Helvetica"
    style.font.size = Pt(11)

    h = doc.add_heading(title, level=0)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)

    for s in sections:
        heading = s.get("heading")
        if heading:
            doc.add_heading(heading, level=1)
        for para in s.get("paragraphs") or []:
            doc.add_paragraph(para)
        for bullet in s.get("bullets") or []:
            doc.add_paragraph(bullet, style="List Bullet")

    fname = f"doc-{uuid.uuid4().hex[:8]}.docx"
    path = os.path.join(OUTPUT_DIR, fname)
    doc.save(path)
    return path
