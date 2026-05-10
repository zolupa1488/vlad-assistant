"""Read incoming Telegram documents into plain text for Claude.

Supported formats: PDF, DOCX, XLSX, CSV, TSV, TXT, MD, JSON, YAML.
Hard-cap output to ~12K characters so it fits into the LLM context window.
"""

from __future__ import annotations

import asyncio
import csv
import os
from pathlib import Path

_MAX_CHARS = 12_000


def _read_pdf(path: str) -> str:
    from pypdf import PdfReader

    reader = PdfReader(path)
    pages = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        pages.append(f"--- page {i + 1} ---\n{text}")
    return "\n\n".join(pages)


def _read_docx(path: str) -> str:
    from docx import Document

    doc = Document(path)
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            parts.append("\t".join(cells))
    return "\n".join(parts)


def _read_xlsx(path: str) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(path, data_only=True, read_only=True)
    parts: list[str] = []
    for sheet in wb.worksheets:
        parts.append(f"=== sheet: {sheet.title} ===")
        for row in sheet.iter_rows(values_only=True):
            line = "\t".join("" if c is None else str(c) for c in row)
            parts.append(line)
    return "\n".join(parts)


def _read_csv(path: str, delim: str) -> str:
    rows: list[str] = []
    with open(path, encoding="utf-8", errors="replace", newline="") as f:
        for row in csv.reader(f, delimiter=delim):
            rows.append("\t".join(row))
    return "\n".join(rows)


def _read_text(path: str) -> str:
    with open(path, encoding="utf-8", errors="replace") as f:
        return f.read()


def _detect_and_read(path: str, name_hint: str | None) -> str:
    name = (name_hint or os.path.basename(path)).lower()
    ext = Path(name).suffix
    if ext == ".pdf":
        return _read_pdf(path)
    if ext == ".docx":
        return _read_docx(path)
    if ext in (".xlsx", ".xlsm"):
        return _read_xlsx(path)
    if ext == ".csv":
        return _read_csv(path, ",")
    if ext == ".tsv":
        return _read_csv(path, "\t")
    if ext in (".txt", ".md", ".log", ".json", ".yaml", ".yml"):
        return _read_text(path)
    # default: try as text
    try:
        return _read_text(path)
    except Exception as e:
        raise ValueError(f"unsupported file type: {ext or 'unknown'} ({e})") from e


async def parse_file(path: str, name_hint: str | None = None) -> str:
    """Extract text content from any supported file format. Truncated to ~12K chars."""
    text = await asyncio.to_thread(_detect_and_read, path, name_hint)
    text = text.strip()
    if len(text) > _MAX_CHARS:
        text = text[:_MAX_CHARS] + f"\n\n…(truncated, total {len(text):,} chars)"
    return text
